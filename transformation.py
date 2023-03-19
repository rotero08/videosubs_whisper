import ffmpy
import whisper
import os
import math
import moviepy.editor as mp
import datetime
import translations
import subprocess
import time
import docx



class Transformations:

    @staticmethod
    def extract_audio(video_file: str):
        root_ext = os.path.splitext(video_file)
        audio_file = f'{video_file.replace(root_ext[1] , "")}.mp3'
        video = mp.VideoFileClip(video_file)
        video.audio.write_audiofile(audio_file)
        return audio_file

    @staticmethod
    def audio_transcribe(audio: str, use_model: str):
        model = whisper.load_model(use_model)
        result = model.transcribe(audio)
        segments = result["segments"]
        subs_file = f'{audio.replace(".mp3" , "")}.srt'
        return [segments, subs_file]
    
    @staticmethod
    def str_to_docx(output_filename, language):
        with open(output_filename, 'r', encoding='utf-8') as f:
            text = f.read()

        # Create a new .docx file and add a paragraph for each line in the .txt file
        doc = docx.Document()
        for line in text.split('\n'):
            doc.add_paragraph(line)

        # Save the .docx file
        doc.save(output_filename.replace(".srt" , ".docx"))
        output_filename = output_filename.replace(".srt" , ".docx")
        return output_filename
    
    @staticmethod
    def docx_to_str(output_filename):
        doc = docx.Document(output_filename)

        with open(f'{output_filename.replace(" es.docx" , "")}.srt', 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                f.write(para.text)
                f.write('\n')
    
    @staticmethod
    def delete_docx(output_filename):
        os.remove(output_filename)
        os.remove(output_filename.replace(" es.docx" , ".docx"))

    @staticmethod
    def text_translate(output_filename, language): #########################change
        #output_filename = Transformations.str_to_docx(output_filename, language)
        #output_filename = translations.translate(output_filename)
        #output_filename = "G:\\1)Descargas\\Universidad Nacional De Bogota\\ICS\\8. Ascent II\\4. Shoot Examples\\Nueva carpeta\\asd\\Biology 1_spanish.docx"
        #Transformations.docx_to_str(output_filename)
        #Transformations.delete_docx(output_filename)
        return
    
    @staticmethod
    def create_srt_file(data, output_filename, language):

        with open(output_filename, 'w', encoding='utf-8') as f:
            for i, item in enumerate(data):
                start = datetime.timedelta(seconds=item['start'])
                end = datetime.timedelta(seconds=item['end'])
                start_srt = Transformations.start_to_srt(start)
                end_srt = Transformations.start_to_srt(end)
                f.write(str(i+1) + '\n')
                f.write(start_srt + ' --> ' + end_srt + '\n')
                f.write(item['text'] + '\n')
                f.write('\n')
        
        if language != 'english':
            Transformations.text_translate(output_filename, language)

    @staticmethod
    def start_to_srt(start):
        hours, remainder = divmod(start.seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        milliseconds = math.floor(start.microseconds / 1000)
        return f"{hours:02d}:{minutes:02d}:{seconds:02d},{milliseconds:03d}"

    
    @staticmethod
    def delete_audio(audio_file):
        os.remove(audio_file)
    
    @staticmethod
    def embedd(video_file: str, subs_files: dict):
        root_ext = os.path.splitext(video_file)
        output_file = f'{video_file.replace(root_ext[1] , "")}_subs.mkv'

        inputs = {video_file: None}
        outputs = {output_file: '-y -c copy -map 0:v -map 0:a '}
        
        i=0
        for dubs, address in subs_files.items():
            inputs[address] = None
            outputs[output_file] += f'-map {i+1}:s:0 -metadata:s:s:{i} language={dubs[:min(len(dubs), 3)]} -metadata:s:s:{i} title={dubs} '
            i += 1
            
        print(inputs, outputs)
        ff = ffmpy.FFmpeg(inputs=inputs, outputs=outputs)
        
        ff.run()
        return
    
    @staticmethod
    def delete_srt(subs_files):
        for file in subs_files.values():
            os.remove(file)

    @staticmethod
    def overwrite_video(video):
        os.remove(video)
        root_ext = os.path.splitext(video)
        video = video.replace(root_ext[1], "_subs.mkv")
        os.rename(video, video.replace("_subs.mkv", ".mkv"))
    
    @staticmethod
    def process_video(parser, video):
        audio = Transformations.extract_audio(video)
        segments, subs_file = Transformations.audio_transcribe(audio, getattr(parser.parse_known_args()[0], "model")[0])
        Transformations.delete_audio(audio)

        subs_files = {}
        for language in getattr(parser.parse_known_args()[0], "subs"):
            address = f'{audio.replace(".mp3" , "")}_{language}.srt'
            subs_files = {**subs_files, **{language: address}}
            Transformations.create_srt_file(segments, address, language)
        
        Transformations.embedd(video, subs_files)
        
        args = parser.parse_args()

        if args.overwrite:
            args.delete = True
        
        if (args.delete):
            Transformations.delete_srt(subs_files)


